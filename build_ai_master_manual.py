from pathlib import Path
from datetime import date
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path(r"C:\Users\Administrator\ai_tutor\AI_Master_技术手册.docx")

NAVY = "0B1D3A"
BLUE = "2E74B5"
CYAN = "1E8C99"
PALE = "E8EEF5"
LIGHT = "F4F6F9"
GRAY = "63708A"
INK = "15233A"
WHITE = "FFFFFF"
GOLD = "7A5A00"
RED = "9B1C1C"

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")

def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)

def add_bottom_border(paragraph, color=BLUE, size="10"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "8")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    p_pr.append(borders)

def set_font(run, size=None, color=None, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic

def set_para(paragraph, before=0, after=6, line=1.25, align=None):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if align is not None:
        paragraph.alignment = align

def add_text(doc, text, size=11, color=INK, bold=False, before=0, after=6, line=1.25, align=None):
    p = doc.add_paragraph()
    set_para(p, before, after, line, align)
    r = p.add_run(text)
    set_font(r, size, color, bold)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(.375 + level * .25)
    p.paragraph_format.first_line_indent = Inches(-.187)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_font(r, 10.5, INK)
    return p

def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(.375)
    p.paragraph_format.first_line_indent = Inches(-.187)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_font(r, 10.5, INK)
    return p

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    sizes = {1: 16, 2: 13, 3: 12}
    colors = {1: BLUE, 2: BLUE, 3: "1F4D78"}
    set_para(p, before={1:16,2:12,3:8}[level], after={1:8,2:6,3:4}[level], line=1.1)
    r = p.add_run(text)
    set_font(r, sizes[level], colors[level], True)
    return p

def add_callout(doc, label, text, fill="F4F6F9", accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    set_para(p, 0, 2, 1.2)
    r = p.add_run(label.upper() + "  ")
    set_font(r, 9, accent, True, name="Consolas")
    r = p.add_run(text)
    set_font(r, 10.5, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)

def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, head in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_shading(cell, PALE)
        p = cell.paragraphs[0]
        set_para(p, 0, 0, 1.1)
        r = p.add_run(head)
        set_font(r, 9, NAVY, True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            set_para(p, 0, 0, 1.15)
            r = p.add_run(str(value))
            set_font(r, 9.5, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table

def add_code(doc, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "102038")
    p = cell.paragraphs[0]
    set_para(p, 0, 0, 1.15)
    r = p.add_run(text)
    set_font(r, 9, "DCEBFF", name="Consolas")
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para(p, 0, 0, 1.0)
    r = p.add_run("AI Master 技术手册 | 内部运维参考 | ")
    set_font(r, 8, GRAY)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    p._p.append(field)

def add_header(section):
    header = section.header
    p = header.paragraphs[0]
    set_para(p, 0, 0, 1.0)
    r = p.add_run("AI MASTER / TECHNICAL OPERATIONS MANUAL")
    set_font(r, 8, GRAY, True, name="Consolas")

def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(.492)
    section.footer_distance = Inches(.492)
    add_header(section); add_footer(section)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"; normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)

    # Cover
    for _ in range(5): doc.add_paragraph()
    kicker = add_text(doc, "TECHNICAL OPERATIONS MANUAL", 10, CYAN, True, 0, 18, 1, WD_ALIGN_PARAGRAPH.CENTER)
    title = add_text(doc, "AI Master\n技术手册", 30, NAVY, True, 0, 8, 1.05, WD_ALIGN_PARAGRAPH.CENTER)
    subtitle = add_text(doc, "本地学习平台、网页访问、账号授权与日常运维指南", 14, "35506E", False, 0, 30, 1.2, WD_ALIGN_PARAGRAPH.CENTER)
    line = doc.add_paragraph(); line.alignment = WD_ALIGN_PARAGRAPH.CENTER; set_para(line, 0, 18, 1)
    run = line.add_run("AI LEARNING PLATFORM  |  FLASK  |  CLOUD TUNNEL  |  LEARNING COACH")
    set_font(run, 9, GOLD, True, name="Consolas")
    meta = add_table(doc, ["文档信息", "内容"], [
        ["版本", "1.0"], ["适用对象", "系统管理员、内容维护者、日常运营人员"],
        ["部署模式", "Windows 本机服务 + Cloudflare Quick Tunnel 临时网页访问"],
        ["更新日期", str(date.today())], ["保密等级", "内部运维资料；不得包含或传播账户密码、API Key、访问令牌"],
    ], [2100, 7260])
    doc.add_page_break()

    add_heading(doc, "1. 文档目的与系统边界", 1)
    add_text(doc, "本手册用于说明 AI Master 的实际运行架构、启动方式、用户授权、网页分享、AI 学习功能、数据维护与故障处理。它面向系统所有者及被授权的维护人员，帮助在不破坏现有学习内容的前提下稳定运行和管理平台。")
    add_callout(doc, "核心原则", "AI Master 当前采用“本机托管 + 临时公网入口”的运行模式。公网访问依赖系统所有者电脑保持在线，Cloudflare Quick Tunnel 链接在隧道重启后会变化。")
    add_heading(doc, "1.1 系统能力概览", 2)
    add_table(doc, ["能力域", "当前能力", "主要入口"], [
        ["学习主界面", "章节航线、学习进度、收藏、错题、练习场、环境声场", "/dashboard"],
        ["知识星海", "10 个知识星系、57 个知识星球、关系连线、章节跳转", "/knowledge-stars"],
        ["沉浸远征", "可拖拽 3D 航线与 AI 学习章节入口", "/static/ai_odyssey.html"],
        ["星辰学习教练", "目标访谈、个人路线、费曼讲解诊断、动态推进", "/static/interview.html"],
        ["授权管理", "创建或审批用户、有效期、AI 配额、封禁、审计", "本机 5100 管理端"],
        ["网页分享", "Cloudflare Quick Tunnel 临时公网 URL", "桌面“网页版（免费试用）”启动器"],
    ], [1800, 5000, 2560])
    add_heading(doc, "1.2 不在当前范围内的能力", 2)
    add_bullet(doc, "Quick Tunnel 不是固定域名服务。电脑重启、cloudflared 进程退出或重新启动网页分享器后，旧链接可能失效。")
    add_bullet(doc, "本机服务不等同于云服务器；同时在线人数与稳定性取决于电脑性能、带宽、网络环境及 AI 服务可用性。")
    add_bullet(doc, "授权管理端默认仅供本机管理员使用，不应通过公网直接暴露。")

    add_heading(doc, "2. 总体架构", 1)
    add_text(doc, "AI Master 由主学习服务、独立授权服务、本地用户数据、静态交互资源与外部 AI 服务组成。对外网页请求通过 Cloudflare 临时隧道转发到本机 5001 端口；授权和 AI 配额仍由本机授权服务控制。")
    add_heading(doc, "2.1 逻辑链路", 2)
    add_code(doc, "访客浏览器\n  -> Cloudflare Quick Tunnel（临时 trycloudflare.com 地址）\n  -> AI Master Flask 主服务 :5001\n       -> 本机授权管理服务 :5100\n       -> 本地数据目录 data/\n       -> DeepSeek AI 服务（按授权配额调用）")
    add_heading(doc, "2.2 主要目录", 2)
    add_table(doc, ["目录或文件", "用途", "维护提示"], [
        ["C:\\Users\\Administrator\\ai_tutor", "AI Master 主项目", "主服务、课程、页面、脚本、可执行打包文件"],
        ["ai_tutor\\app.py", "主 Flask 服务", "登录、课程、学习数据、AI 接口、静态页面保护"],
        ["ai_tutor\\templates", "动态 HTML 模板", "仪表盘、章节、登录等服务器渲染页面"],
        ["ai_tutor\\static", "静态页面与资源", "CG、知识星海、沉浸远征、学习教练、CSS、JS"],
        ["ai_tutor\\data", "主服务可写数据", "用户进度、错题、收藏、学习教练路线、运行日志"],
        ["C:\\Users\\Administrator\\ai_tutor_license_admin", "独立授权管理项目", "用户审批、期限、AI 配额、后台审计"],
    ], [2850, 3000, 3510])
    add_callout(doc, "数据原则", "课程与静态内容属于项目资源；用户进度、授权、学习教练路线属于运行数据。升级前应优先备份两个 data 目录。", "FFF8E8", GOLD)

    add_heading(doc, "3. 本机启动与日常操作", 1)
    add_text(doc, "桌面保留三类启动入口。正常情况下不需要在命令行手工启动 Python 服务。启动脚本会检测端口、缺失服务并按需启动。")
    add_table(doc, ["桌面启动器", "用途", "启动后结果"], [
        ["启动AI Master.bat", "本机学习与测试", "启动授权服务与 AI Master，打开 http://127.0.0.1:5001"],
        ["启动 AI Master 授权管理.bat", "管理用户权限与 AI 配额", "启动授权管理服务，打开 http://127.0.0.1:5100/login"],
        ["启动 AI Master 网页版（免费试用）.bat", "生成临时公网访问链接", "启动本机服务和 Quick Tunnel，并更新桌面链接文件"],
    ], [3000, 3000, 3360])
    add_heading(doc, "3.1 推荐启动顺序", 2)
    add_number(doc, "日常本机使用：双击“启动AI Master.bat”，在浏览器登录后进入仪表盘。")
    add_number(doc, "需要授权新用户：双击“启动 AI Master 授权管理.bat”，以管理员身份登录管理端并配置用户期限与 AI 配额。")
    add_number(doc, "需要外部测试：双击“启动 AI Master 网页版（免费试用）.bat”，等待桌面的“AI Master Web Link.txt”更新，再发送其中的新链接。")
    add_callout(doc, "重要", "临时网页链接不是永久域名。每次重新启动网页分享器，都必须以桌面链接文件中的最新地址为准。")

    add_heading(doc, "4. 用户、授权与 AI 配额", 1)
    add_text(doc, "AI Master 将学习账户与授权账户分层处理。外部用户通过网页注册并登录；只有被管理员批准、处于有效期内且仍有 AI 配额的账户，才可以继续使用受保护页面与 AI 功能。")
    add_heading(doc, "4.1 用户授权流程", 2)
    add_number(doc, "用户获取管理员提供的网页链接，在登录页提交注册申请。")
    add_number(doc, "管理员在授权管理端找到该用户，设置有效天数与 AI 使用上限，然后授予访问权限。")
    add_number(doc, "用户使用注册账号登录 AI Master。登录会获取一个有效期受控的授权会话。")
    add_number(doc, "每次 AI 调用均由授权服务检查状态、有效期与使用额度；用尽时需要管理员增加额度或重置配额。")
    add_heading(doc, "4.2 管理端可执行动作", 2)
    add_table(doc, ["操作", "作用", "建议"], [
        ["创建 / 批准用户", "将待审核账号置为可用状态", "确认用户名与使用期限后再批准"],
        ["设置有效天数", "控制用户可登录期限", "短期体验账户使用较短期限，正式学员按课程周期设置"],
        ["设置 AI 上限", "控制 DeepSeek 等 AI 功能调用次数", "为试用用户设置较小额度；观察使用后再调整"],
        ["重置配额", "将已用 AI 次数清零", "适用于续费、补发或测试"],
        ["封禁 / 拒绝", "立即阻止账户继续访问", "异常访问或不再授权时使用；现有令牌会失效"],
        ["审计日志", "记录注册、授权、登录和管理动作", "定期检查异常账号和配额变动"],
    ], [1900, 3400, 4060])
    add_callout(doc, "安全提醒", "管理员密码、DeepSeek API Key、授权令牌和 session secret 均属于敏感信息。不要写入教程截图、群聊、公开文档或压缩包说明。")

    add_heading(doc, "5. 网页访问与 Cloudflare Quick Tunnel", 1)
    add_text(doc, "当前公网访问方式使用 Cloudflare Quick Tunnel。该方式适合少量体验用户和临时展示，优点是无需购买域名；缺点是 URL 会变化且可用性受本机网络、电脑在线状态和部分访问网络策略影响。")
    add_heading(doc, "5.1 正确分享方式", 2)
    add_number(doc, "启动“AI Master 网页版（免费试用）”。")
    add_number(doc, "打开桌面“AI Master Web Link.txt”，复制其中以 https:// 开头的完整地址。")
    add_number(doc, "将最新地址发送给测试用户，并说明需要先注册、等待授权后登录。")
    add_number(doc, "电脑保持开机、联网，AI Master 服务和 cloudflared 进程不能被结束。")
    add_heading(doc, "5.2 失效判断与处理", 2)
    add_table(doc, ["现象", "常见原因", "处理方式"], [
        ["旧链接无法打开", "隧道已重启或 cloudflared 已退出", "重新运行网页版启动器，发送新的链接"],
        ["本机可开，外部用户打不开", "外部网络拦截 trycloudflare.com、临时链路波动或电脑离线", "先确认本机服务正常；让用户换网络测试；必要时重新生成链接"],
        ["登录后马上退出", "授权过期、用户被封禁、授权服务不可达", "检查授权管理端用户状态、有效期与本机 5100 服务"],
        ["页面能打开但 AI 不可用", "AI 配额耗尽或外部 AI 服务不可达", "查看授权额度与 AI 服务配置；学习教练会启用本地保底路线"],
    ], [2400, 3500, 3460])
    add_callout(doc, "长期方案", "当需要稳定固定网址、更多并发或不依赖个人电脑时，应迁移到已购域名的 Cloudflare Named Tunnel 或云服务器部署。")

    add_heading(doc, "6. 学习产品与互动页面", 1)
    add_text(doc, "AI Master 不只是课程列表。核心体验由仪表盘、知识星海、章节探索、沉浸远征和个人学习教练共同构成。各页面通过登录会话保护，未登录访问静态学习页面会被重定向到登录页。")
    add_table(doc, ["模块", "面向用户的价值", "运行特性"], [
        ["仪表盘", "统一查看章节、进度、工具和星域入口", "章节状态由用户进度与探索模式决定"],
        ["知识星海", "在可旋转星系中查看知识节点与关系网络", "Three.js 本地资源；10 星系、57 节点；点击进入学习"],
        ["章节探索", "在单章中沿知识航标学习、练习与复盘", "保留既有内容、收藏、错题和章节 AI 辅导"],
        ["沉浸远征", "用 3D 航线进入 Transformer、Prompt、Agent、RAG 等章节", "独立页面；支持拖拽、滚动推进和性能降级"],
        ["星辰学习教练", "由学习目标驱动的个人路线与费曼输出", "路线按用户保存，可从不同设备继续执行"],
    ], [1800, 4100, 3460])
    add_heading(doc, "6.1 星辰学习教练工作流", 2)
    add_code(doc, "目标访谈 -> 结合课程、进度与错题生成路线\n        -> 进入指定章节或知识星球\n        -> 用户以费曼法讲解\n        -> AI 诊断理解缺口、给出追问和微课\n        -> 达标后点亮下一里程碑；未达标则补弱后再讲")
    add_bullet(doc, "优先使用现有 DeepSeek 调用链路生成个性化路线和讲解反馈；所有 AI 调用仍受账户授权和 AI 额度约束。")
    add_bullet(doc, "AI 服务暂时不可用时，系统会按学习目标和课程目录生成保底路线，并用规则化费曼检查保持页面可用。")
    add_bullet(doc, "学习教练路线与讲解记录保存在用户数据中，不与其他用户共享。")

    add_heading(doc, "7. 数据、备份与恢复", 1)
    add_text(doc, "运行数据主要保存于主项目和授权管理项目各自的 data 目录。备份时应在服务停止后进行，或至少确保写入动作较少，以避免复制到半写入状态。")
    add_table(doc, ["数据位置", "内容", "备份优先级"], [
        ["ai_tutor\\data\\users.json", "本地管理员及本机学习数据：进度、错题、收藏、JJ 记录、学习教练路线", "高"],
        ["ai_tutor_license_admin\\data\\licenses.db", "授权用户、状态、有效期、AI 配额、审计日志", "最高"],
        ["ai_tutor\\data\\.session_secret", "会话签名密钥；替换会使所有用户重新登录", "高"],
        ["ai_tutor_license_admin\\data\\.admin_secret", "授权管理后台会话密钥", "高"],
        ["ai_tutor\\data\\chapter_*.json", "课程章节数据与内容", "高"],
        ["ai_tutor\\static 与 templates", "页面、交互、动画、CSS 和 JavaScript", "高"],
    ], [3000, 4300, 2060])
    add_heading(doc, "7.1 建议备份流程", 2)
    add_number(doc, "关闭网页分享和本机服务，或确认没有人在进行写入操作。")
    add_number(doc, "复制 ai_tutor\\data、ai_tutor_license_admin\\data、static 和 templates 到以日期命名的备份目录。")
    add_number(doc, "不要把包含密钥、数据库或用户信息的备份发送给普通学员。")
    add_number(doc, "恢复时先备份当前目录，再覆盖恢复数据，最后依次启动授权管理端与 AI Master。")

    add_heading(doc, "8. 安全与健壮性要求", 1)
    add_table(doc, ["控制项", "当前做法", "管理员要求"], [
        ["页面访问", "登录会话保护受保护的静态学习页面与 API", "不要把内部管理端公开到临时公网"],
        ["用户授权", "授权服务校验状态、到期时间与令牌", "仅批准确认过的用户；离开课程后及时封禁或到期"],
        ["AI 配额", "每次 AI 调用在授权服务侧计数", "合理设置试用上限，防止无控制消耗"],
        ["请求保护", "主服务对 API 进行限流，并配置基础安全响应头", "异常高频访问时先检查日志和授权记录"],
        ["静态资源", "Three.js 等核心互动依赖采用本地文件", "避免依赖不可控 CDN 导致外网用户页面空白"],
        ["功能降级", "WebGL 与学习教练均具备失败回退", "遇到第三方服务波动时优先保持课程和基础路线可用"],
    ], [1800, 4400, 3160])
    add_callout(doc, "禁止事项", "不得在代码仓库、截图、手册、公开网页或群聊中保留明文 API Key、管理员密码或用户令牌。发现已泄露密钥时，应立即在服务端轮换。", "FDEBEC", RED)

    add_heading(doc, "9. 常见故障排查", 1)
    add_heading(doc, "9.1 页面显示异常或只剩裸文字", 2)
    add_bullet(doc, "确认用户从 AI Master 登录后进入页面，而不是绕过服务直接用 file:// 打开需要服务器资源的页面。")
    add_bullet(doc, "检查浏览器是否加载到 CSS、JavaScript 和本地 Three.js 文件；刷新页面或清除强缓存后重试。")
    add_bullet(doc, "若 WebGL 不可用，知识星海和沉浸远征应降级或提示错误；可尝试更新浏览器、显卡驱动，或切换到轻量画质。")
    add_heading(doc, "9.2 授权、登录或注册失败", 2)
    add_bullet(doc, "确认 5100 授权管理服务正在本机监听，并从授权管理端检查用户状态、到期时间与 AI 上限。")
    add_bullet(doc, "确认 5001 主服务正在运行。若网页分享模式使用中，两个服务都必须持续运行。")
    add_bullet(doc, "外部用户注册后需要管理员批准；注册成功不等于立即具备登录权限。")
    add_heading(doc, "9.3 AI 学习教练未生成个性化结果", 2)
    add_bullet(doc, "先检查该账号的 AI 额度是否耗尽，以及授权管理端是否配置了 AI 服务。")
    add_bullet(doc, "若外部 AI 网络暂时失败，学习教练会生成基于课程目录的保底路线。等待服务恢复后，用户可点击“重新校准路线”获得更个性化结果。")
    add_heading(doc, "9.4 端口与进程检查", 2)
    add_code(doc, "PowerShell:\nGet-NetTCPConnection -LocalPort 5001 -State Listen\nGet-NetTCPConnection -LocalPort 5100 -State Listen\nGet-Content C:\\Users\\Administrator\\ai_tutor\\data\\web-share-tunnel.log -Tail 30")

    add_heading(doc, "10. 更新、发布与验收清单", 1)
    add_text(doc, "任何页面、课程或授权逻辑更新后，都建议先在本机登录环境验收，再开启网页分享给外部用户。对于涉及主服务 app.py 的更新，必须重启 AI Master 服务使新接口生效。")
    add_heading(doc, "10.1 更新后验收", 2)
    for item in [
        "本机登录、仪表盘、章节页面和知识星海均可正常打开。",
        "沉浸远征与星辰学习教练分别作为独立入口存在，未覆盖 3D 知识星海。",
        "学习教练可读取已有路线；新用户可完成目标访谈并看到学习星图。",
        "用户完成费曼讲解后能得到反馈；达标后可推进下一里程碑。",
        "授权管理端可以查看、批准、封禁用户并调整 AI 配额。",
        "运行网页版启动器后，桌面链接文件已更新；用外部网络至少完成一次登录和章节跳转测试。",
        "主服务、授权服务日志中没有持续报错；确认敏感信息未被写入页面或文档。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "10.2 面向测试用户的简短说明", 2)
    add_callout(doc, "可直接发送", "请使用最新网页链接访问 AI Master。首次使用请先注册账号，并将用户名发给管理员等待授权。授权后可登录学习；如链接失效，请向管理员索取最新地址。", "EAF7F4", CYAN)

    add_heading(doc, "附录 A：关键端点与页面索引", 1)
    add_table(doc, ["路径", "功能", "访问控制"], [
        ["/login", "用户登录与注册申请", "公开入口"],
        ["/dashboard", "主学习仪表盘", "登录及授权有效"],
        ["/knowledge-stars", "3D 知识星海", "登录及授权有效"],
        ["/static/ai_odyssey.html", "沉浸远征", "登录及授权有效"],
        ["/static/interview.html", "星辰学习教练", "登录及授权有效"],
        ["/api/knowledge-universe", "知识星海数据", "登录及授权有效"],
        ["/api/learning-coach/state", "读取个人学习路线", "登录及授权有效"],
        ["/api/learning-coach/plan", "生成或重建个人路线", "登录及授权有效；消耗 AI 配额或使用保底路线"],
        ["/api/learning-coach/feynman", "费曼讲解诊断", "登录及授权有效；消耗 AI 配额或使用保底反馈"],
        ["/api/learning-coach/advance", "推进当前里程碑", "登录及授权有效；需先通过讲解检验"],
    ], [3100, 4100, 2160])
    add_heading(doc, "附录 B：维护联系人与记录", 1)
    add_text(doc, "建议在每次较大更新后记录：更新日期、修改模块、备份位置、是否重启 5001/5100 服务、是否重新生成公网链接、外网验收结果。此记录可另存为运维日志。")
    add_callout(doc, "结束语", "稳定运营的关键不是一直生成新页面，而是把“启动、授权、备份、外网验证、故障恢复”变成每次更新后的固定流程。", "E8EEF5", BLUE)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "AI Master 技术手册"
    doc.core_properties.subject = "AI Master 运维、授权与网页部署指南"
    doc.core_properties.author = "AI Master"
    doc.save(OUT)
    print(OUT)

if __name__ == "__main__":
    build_doc()
